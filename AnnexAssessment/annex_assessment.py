import os
import io
import re
import json
import pandas as pd
from datetime import datetime
import streamlit as st
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.oauth2 import service_account
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, ValidationError
from typing import Optional, List

###############################################################################
# 0. DATA MODEL & OPENAI CLASSIFICATION
###############################################################################

class FileClassification(BaseModel):
    file_id: str
    file_name: str
    is_assessment_plan: bool = False
    assessment_type: Optional[str] = None
    is_question_paper: bool = False
    is_answer_paper: bool = False
    version: Optional[str] = None


def classify_files_with_openai(file_metadata: List[dict]) -> List[FileClassification]:
    """
    Uses OpenAI to classify files into assessment plan, question paper, or answer paper.
    Returns a list of FileClassification objects (only for relevant files).
    """
    from openai import OpenAI
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    # Prepare file metadata for OpenAI
    file_info = "\n".join([f"{file['id']} - {file['name']}" for file in file_metadata])
    system_message = (
        "You are an AI assistant tasked with analyzing file names related to WSQ assessments. "
        "For each file, identify:\n"
        "1. Whether it is an assessment plan, question paper, or answer paper.\n"
        "   - Recognize variations like 'AP_' prefixes for assessment plans or '(Draft)' in filenames.\n"
        "2. For question or answer papers, identify the assessment type (e.g., WA (SAQ), PP, CS, Oral Questioning).\n"
        "   - Include incomplete or draft versions, noting them explicitly.\n"
        "3. Extract the version (e.g., v2.1, v1.5) if available.\n"
        "4. If a file cannot be classified, mark it as irrelevant.\n"
        "\n"
        "Return only a valid JSON array with no additional text or explanation, using this format:\n"
        "[{\n"
        "  \"file_id\": \"\",\n"
        "  \"file_name\": \"\",\n"
        "  \"is_assessment_plan\": false,\n"
        "  \"assessment_type\": null,\n"
        "  \"is_question_paper\": false,\n"
        "  \"is_answer_paper\": false,\n"
        "  \"version\": \"\"\n"
        "}]\n"
        "\n"
        "If no valid classifications are found, return an empty JSON array: []"
    )

    user_message = (
        "Below is a list of files from a WSQ course folder. "
        "Classify each file as an assessment plan, question paper, or answer paper. "
        "Recognize drafts (e.g., '(Draft)') and variations in naming. "
        "For unclassifiable files, mark them as irrelevant. "
        "Return the results in valid JSON format only.\n\n"
        "Files:\n"
        f"{file_info}\n\n"
        "Provide the classification in JSON format only."
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_message},
        ],
        max_tokens=1000,
        temperature=0.3,
    )

    content = response.choices[0].message.content.strip()
    try:
        # Attempt to extract JSON content
        json_start = content.find("[")
        json_end = content.rfind("]") + 1
        if json_start != -1 and json_end != -1:
            json_content = content[json_start:json_end]
        else:
            raise ValueError("No JSON array found in response")

        json_content = json_content.replace("'", '"')  # ensure valid JSON
        raw_list = json.loads(json_content)

        # Filter only relevant files
        filtered_list = [
            item for item in raw_list
            if item["is_assessment_plan"] or item["is_question_paper"] or item["is_answer_paper"]
        ]

        return [FileClassification(**item) for item in filtered_list]

    except (json.JSONDecodeError, ValidationError, ValueError) as e:
        print("Error parsing OpenAI response:", e)
        print("OpenAI response content was:", content)
        return []


###############################################################################
# 1. HELPER FUNCTIONS
###############################################################################

def authenticate():
    """
    Authenticate with Google using credentials from Streamlit secrets.
    """
    try:
        creds = service_account.Credentials.from_service_account_info(
            st.secrets["GOOGLE_API_CREDS"]
        )
        return creds
    except Exception as e:
        print(f"Error during authentication: {e}")
        return None


def download_file(file_id, file_name, drive_service, download_dir="./downloads"):
    """
    Downloads a file (Google Doc or Word .docx) from Google Drive.
    Returns the local path to the downloaded file, or None if skipped.
    """
    if not os.path.exists(download_dir):
        os.makedirs(download_dir, exist_ok=True)

    file_info = drive_service.files().get(fileId=file_id, fields="mimeType").execute()
    mime_type = file_info.get("mimeType")

    if mime_type == "application/vnd.google-apps.document":
        export_mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        request = drive_service.files().export_media(fileId=file_id, mimeType=export_mime_type)
        base_name, _ = os.path.splitext(file_name)
        file_name = base_name + ".docx"
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        request = drive_service.files().get_media(fileId=file_id)
    else:
        print(f"Skipping file (not .docx or Google Doc): {file_name}")
        return None

    file_path = os.path.join(download_dir, file_name)
    with io.BytesIO() as fh:
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)
        with open(file_path, "wb") as f:
            f.write(fh.read())

    return file_path


def parse_version(version_str: Optional[str]) -> tuple:
    """
    Extracts a (major, minor) tuple from a version string like 'v2.1' or 'v1'.
    Returns (0,0) for invalid or missing versions.
    """
    if version_str:
        match = re.match(r"v(\d+)(\.\d+)?", version_str.lower())
        if match:
            major = int(match.group(1))
            minor = int(match.group(2).lstrip(".")) if match.group(2) else 0
            return (major, minor)
    return (0, 0)


def select_latest_version(file_classifications: List[FileClassification]) -> Optional[FileClassification]:
    """
    Given a list of FileClassification objects, returns the one with the highest 'version'.
    """
    sorted_files = sorted(
        file_classifications,
        key=lambda f: parse_version(f.version),
        reverse=True
    )
    return sorted_files[0] if sorted_files else None


def select_latest_assessment_plan(file_classifications: List[FileClassification]) -> Optional[FileClassification]:
    """
    Among the classifications, picks the latest version of an assessment plan.
    """
    plans = [f for f in file_classifications if f.is_assessment_plan]
    return select_latest_version(plans)


def build_method_data(file_classifications: List[FileClassification], abbreviations: List[str]) -> dict:
    """
    Builds a dictionary of {abbr: {"question": {...}, "answer": {...}}}
    picking the latest version for question and answer papers.
    """
    method_data = {}

    for abbr in abbreviations:
        question_files = [
            f for f in file_classifications
            if f.assessment_type == abbr and f.is_question_paper
        ]
        answer_files = [
            f for f in file_classifications
            if f.assessment_type == abbr and f.is_answer_paper
        ]

        latest_question = select_latest_version(question_files)
        latest_answer = select_latest_version(answer_files)

        if latest_question or latest_answer:
            method_data[abbr] = {
                "question": {
                    "id": latest_question.file_id,
                    "name": latest_question.file_name
                } if latest_question else None,
                "answer": {
                    "id": latest_answer.file_id,
                    "name": latest_answer.file_name
                } if latest_answer else None,
            }

    return method_data


def delete_irrelevant_files(download_dir="./downloads", keep_filename=None):
    """
    Delete all files in 'download_dir' except the 'keep_filename'.
    If 'keep_filename' is None, everything is removed.
    """
    for file_name in os.listdir(download_dir):
        file_path = os.path.join(download_dir, file_name)
        if keep_filename is not None and file_name == keep_filename:
            continue
        try:
            os.remove(file_path)
            print(f"Deleted local file: {file_path}")
        except PermissionError:
            print(f"Could not delete (in use): {file_path}")


###############################################################################
# 2. PROCESS COURSE FOLDER (WITH OPENAI CLASSIFICATION)
###############################################################################

def process_course_folder(course_folder_id, drive_service, abbreviations):
    """
    1) Looks for 'Assessment Plan' and 'Assessment' subfolders.
    2) Classifies files in both folders using OpenAI.
    3) Selects the latest assessment plan.
    4) Builds the method_data dictionary for Q&A docs by abbreviations.
    Returns {"assessment_plan": {...}, "method_data": {...}} or None if no plan found.
    """
    # Retrieve subfolders
    subfolders = drive_service.files().list(
        q=f"'{course_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'"
    ).execute().get('files', [])

    # Find "Assessment Plan" folder (case-insensitive check)
    assessment_plan_folder = next(
        (f for f in subfolders if f['name'].strip().lower() == 'assessment plan'),
        None
    )

    plan_files = []
    if assessment_plan_folder:
        plan_files = drive_service.files().list(
            q=(
                f"'{assessment_plan_folder['id']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
    else:
        print(f"No 'Assessment Plan' folder found in {course_folder_id}.")

    if plan_files:
        plan_classifications = classify_files_with_openai(plan_files)
        assessment_plan = select_latest_assessment_plan(plan_classifications)
    else:
        print(f"No files found in Assessment Plan folder for {course_folder_id}.")
        assessment_plan = None

    # If no plan found, check all files in the course folder for a possible misclassified plan
    if not assessment_plan:
        print(f"No valid assessment plan found for {course_folder_id}. Checking further...")
        all_course_files = drive_service.files().list(
            q=(
                f"'{course_folder_id}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
        all_classifications = classify_files_with_openai(all_course_files)
        assessment_plan = select_latest_assessment_plan(all_classifications)

    if not assessment_plan:
        print(f"No valid assessment plan could be identified for folder ID {course_folder_id}.")
        return None

    # Find "Assessment" folder (case-insensitive check)
    assessment_folder = next(
        (f for f in subfolders if f['name'].strip().lower() == 'assessment'),
        None
    )

    if assessment_folder:
        assessment_files = drive_service.files().list(
            q=(
                f"'{assessment_folder['id']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            )
        ).execute().get('files', [])
    else:
        print(f"No 'Assessment' folder found for {course_folder_id}.")
        assessment_files = []

    if assessment_files:
        assessment_classifications = classify_files_with_openai(assessment_files)
        method_data = build_method_data(assessment_classifications, abbreviations)
    else:
        method_data = {}

    return {
        "assessment_plan": {
            "id": assessment_plan.file_id,
            "name": assessment_plan.file_name
        },
        "method_data": method_data,
    }


###############################################################################
# 3. MERGING DOCUMENTS INTO ANNEX & VERSION UPDATES
###############################################################################
from docx.enum.table import WD_TABLE_ALIGNMENT

def insert_centered_header(doc, text, annex_label):
    """
    Inserts a right-aligned paragraph with 'Annex D' above a centered header.
    The header is centered both vertically and horizontally using a single-cell table.
    A page break is added after the table.
    """

    # Insert the centered header using a single-cell table...
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # or CENTER, RIGHT, etc.

    # Vertical alignment
    table_cell = table.cell(0, 0)
    table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Row height (10 inches as an example)
    row = table.rows[0]._tr
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '12000')  # 12000 twips ~ 8.3 inches; adjust as needed
    rowPr = row.get_or_add_trPr()
    rowPr.append(trHeight)

    # Add the centered text to the table
    paragraph = table_cell.paragraphs[0]
    run = paragraph.add_run(f"{annex_label}:\n{text}")
    run.bold = True
    run.font.size = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break after the table
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def insert_answers_under_heading(plan_path, heading_map, method_data):
    """
    Inserts question and answer papers into the annex of the plan document.
    - heading_map: dict that maps some textual heading to the abbreviation.
    - method_data: output from build_method_data (dict of abbr -> question/answer).
    Returns (updated_doc_path, changes_made).
    """
    base_doc = Document(plan_path)
    composer = Composer(base_doc)
    changes_made = False
    annex_index = 0

    for heading_text, abbr in heading_map.items():
        if abbr in method_data:
            files = method_data[abbr]
            q_file = files.get('question')
            a_file = files.get('answer')

            # If there's a question doc
            if q_file and 'local_path' in q_file:

                annex_label = get_annex_label(annex_index)  # e.g. "Annex A"
                annex_index += 1
                
                temp_doc = Document()
                insert_centered_header(temp_doc, f"QUESTION PAPER OF {abbr} ASSESSMENT", annex_label)
                composer.append(temp_doc)

                question_doc = Document(q_file['local_path'])
                composer.append(question_doc)
                changes_made = True

            # If there's an answer doc
            if a_file and 'local_path' in a_file:

                annex_label = get_annex_label(annex_index)  # e.g. "Annex A"
                annex_index += 1

                temp_doc = Document()
                insert_centered_header(temp_doc, f"SUGGESTED ANSWER TO {abbr} ASSESSMENT QUESTIONS", annex_label)
                composer.append(temp_doc)

                answer_doc = Document(a_file['local_path'])
                composer.append(answer_doc)
                changes_made = True

    if changes_made:
        updated_path = plan_path.replace(".docx", "_Answers_Only.docx")
        composer.save(updated_path)

        # Auto-fit tables in the merged doc
        updated_doc = Document(updated_path)
        updated_doc.save(updated_path)

        return updated_path, True
    else:
        print("No Q&A appended to annex.")
        return plan_path, False


def update_cover_page_version(doc_path):
    """
    Increments the 'Version X' text on the cover page, if found, to 'Version X+1.0'.
    """
    doc = Document(doc_path)
    updated = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Version"):
            # Attempt to parse old version
            try:
                current_version = float(text.split()[1])
                new_version = f"Version {int(current_version) + 1}.0"
            except (IndexError, ValueError):
                raise ValueError("Invalid version format on cover page.")

            paragraph.clear()
            run = paragraph.add_run(new_version)
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            print(f"Updated cover page to: {new_version}")
            updated = True
            break

    if not updated:
        raise ValueError("No 'Version' text found on the cover page.")

    updated_doc_path = doc_path.replace(".docx", "_Updated.docx")
    doc.save(updated_doc_path)
    print(f"Cover page version updated => {updated_doc_path}")
    return updated_doc_path


def update_version_number(last_version_str):
    """
    Converts a last_version_str like '1.0' or '2.1' to the next major version: '2.0' or '3.0'.
    """
    try:
        last_version = float(last_version_str)
        next_version = int(last_version) + 1
        return f"{next_version}.0"
    except ValueError:
        return "2.0"

def get_annex_label(index: int) -> str:
    """
    Returns 'Annex A' for index=0, 'Annex B' for index=1, etc.
    If index goes beyond 25, you may need to handle it differently (e.g., AA).
    """
    letter = chr(ord("A") + index)
    return f"Annex {letter}"

def update_version_control_record(doc_path, changes, developer="Tertiary Infotech"):
    """
    Appends a new row in the first table to track version updates.
    Expects the first table to be the Version Control table with columns:
    [Version, Effective Date, Changes, Developer].
    """
    doc = Document(doc_path)

    if not doc.tables:
        print("No tables found. Can't update version control record.")
        return

    version_table = doc.tables[0]  # assume first table is version control
    # The last version in the first column
    last_version_str = version_table.rows[-1].cells[0].text.strip()
    next_version_str = update_version_number(last_version_str)

    eff_date = datetime.now().strftime("%d %b %Y")
    row_cells = version_table.add_row().cells
    row_cells[0].text = next_version_str
    row_cells[1].text = eff_date
    row_cells[2].text = changes
    row_cells[3].text = developer

    # Center/align each column
    for idx, cell in enumerate(row_cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            if idx == 2:  # "Changes" column left-aligned
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(doc_path)
    print(f"Version control record updated => {doc_path}")


def bump_filename_version(doc_path):
    """
    Bumps the numeric version in the filename itself:
      e.g. 'xxx_v2.docx' => 'xxx_v3.docx'
    For fractional versions, it increments the float by +1.0 (v2.1 => v3.1).
    If no version is found, doc_path is unchanged.
    """
    base, ext = os.path.splitext(doc_path)
    pattern = r'(?:-|_)v(\d+(\.\d+)*)'

    match = re.search(pattern, base, re.IGNORECASE)
    if not match:
        return doc_path  # no 'vXX' in filename => do nothing

    old_version_str = match.group(1)
    try:
        old_ver_float = float(old_version_str)
        new_ver_float = old_ver_float + 1.0
        # Keep the minor if it existed
        new_version_str = f"{int(new_ver_float)}" if old_ver_float.is_integer() else f"{new_ver_float}"
    except ValueError:
        new_version_str = "2"  # fallback

    def replacement(m):
        prefix = m.group(0)[:-len(old_version_str)]
        return prefix + new_version_str

    new_base = re.sub(pattern, replacement, base, flags=re.IGNORECASE)
    new_doc_path = f"{new_base}{ext}"

    if os.path.exists(doc_path):
        os.rename(doc_path, new_doc_path)
        print(f"Renamed {doc_path} => {new_doc_path}")

    return new_doc_path


def upload_updated_doc(drive_service, file_id, local_doc_path, original_filename):
    """
    Upload the updated doc to Google Drive, removing any '_Answers_Only' or '_Updated'
    from the final name.
    """
    base_name, ext = os.path.splitext(original_filename)
    # Remove suffixes for cleanliness
    base_name = re.sub(r'_Answers_Only', '', base_name, flags=re.IGNORECASE)
    base_name = re.sub(r'_Updated', '', base_name, flags=re.IGNORECASE)
    new_filename = base_name + ext

    media_body = MediaFileUpload(
        local_doc_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True
    )

    updated_file = drive_service.files().update(
        fileId=file_id,
        media_body=media_body,
        body={"name": new_filename}
    ).execute()

    print(f"Updated file in Google Drive: {updated_file.get('name')} (ID: {file_id})")
    return updated_file

###############################################################################
# 4. TRACKING PROCESSES FUNCTION
###############################################################################

def track_edited_assessment_plan(course_title, excel_file="edited_assessment_plans.xlsx"):
    """
    Tracks edited assessment plans in an Excel file.
    Appends a new entry with the course title using pandas.concat.
    """
    # Check if the Excel file exists
    if not os.path.exists(excel_file):
        # Create a new DataFrame with headers if the file doesn't exist
        df = pd.DataFrame(columns=["Course Title"])
    else:
        # Load the existing Excel file
        df = pd.read_excel(excel_file)

    # Create a new DataFrame for the new entry
    new_entry = pd.DataFrame([[course_title]], columns=["Course Title"])

    # Concatenate the new entry with the existing DataFrame
    df = pd.concat([df, new_entry], ignore_index=True)

    # Save the updated DataFrame back to the Excel file
    df.to_excel(excel_file, index=False)
    print(f"Updated tracking file: {excel_file} with new entry: {course_title}")


###############################################################################
# 4. MAIN FUNCTION
###############################################################################

def process_course_folder_direct(course_folder_id, drive_service, abbreviations):
    """
    Directly processes the 'Assessment Plan' and 'Assessment' folders from the specified course folder ID.
    Ensures folder names are stripped and case-insensitive.
    """
    # Define the target folder names
    target_folders = {"assessment plan": None, "assessment": None}

    # Retrieve all subfolders in the course folder
    subfolders = drive_service.files().list(
        q=f"'{course_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'",
        fields="files(id, name)"
    ).execute().get("files", [])

    # Match and map the folder names
    for subfolder in subfolders:
        folder_name = subfolder["name"].strip().lower()  # Normalize folder name
        if folder_name in target_folders:
            target_folders[folder_name] = subfolder["id"]

    # Retrieve 'Assessment Plan' files
    assessment_plan = None
    if target_folders["assessment plan"]:
        plan_files = drive_service.files().list(
            q=(
                f"'{target_folders['assessment plan']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            ),
            fields="files(id, name)"
        ).execute().get("files", [])

        if plan_files:
            plan_classifications = classify_files_with_openai(plan_files)
            assessment_plan = select_latest_assessment_plan(plan_classifications)
        else:
            print(f"No files found in 'Assessment Plan' folder for course folder ID {course_folder_id}.")

    # Retrieve 'Assessment' files
    method_data = {}
    if target_folders["assessment"]:
        assessment_files = drive_service.files().list(
            q=(
                f"'{target_folders['assessment']}' in parents and "
                "(mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' or "
                "mimeType='application/vnd.google-apps.document')"
            ),
            fields="files(id, name)"
        ).execute().get("files", [])

        if assessment_files:
            assessment_classifications = classify_files_with_openai(assessment_files)
            method_data = build_method_data(assessment_classifications, abbreviations)

    # Return the processed data
    if not assessment_plan:
        print(f"No valid assessment plan identified for course folder ID {course_folder_id}.")
        return None

    return {
        "assessment_plan": {
            "id": assessment_plan.file_id,
            "name": assessment_plan.file_name
        },
        "method_data": method_data,
    }

def app():
    """
    Streamlit app to process a course folder by its name, integrating assessment questions
    and answers into the annex of the assessment plan document.
    """
    st.title("📄 Integrate Assessment to Annex of AP")
    st.header("Instructions:")
    st.markdown("""
    Enter the exact course folder name from Google Drive. The app will integrate assessment questions 
    and answers into the annex of the assessment plan document.
    """)
    st.markdown("""
    #### 📂 File Organization and Naming Instructions

    - **`Assessment Plan`** folder:  
    Place the **assessment plan file** here.  
    Example: `Assessment Plan_TGS-[Course Code] - [Course Title]_vX.docx`

    - **`Assessment`** folder:  
    Place **question and answer files** here.  
    Examples:  
    - `WA (SAQ) - [Course Title] - vX.docx`  
    - `Answer to WA (SAQ) - [Course Title] - vX.docx`

    """)


    # Abbreviations and heading map
    abbreviations = ["WA (SAQ)", "PP", "CS", "RP", "Oral Questioning"]
    heading_map = {
        "Assessment Questions and Answers for WA(SAQ)": "WA (SAQ)",
        "Assessment Questions and Practical Performance": "PP",
        "Assessment Questions and Case Study": "CS",
        "Assessment Questions and Oral Questioning (OQ)": "Oral Questioning",
    }

    # Authenticate with Google Drive
    with st.spinner("Authenticating with Google Drive..."):
        creds = authenticate()
    if not creds:
        st.error("Authentication failed. Please check your credentials.")
        return

    drive_service = build("drive", "v3", credentials=creds)

    # Input: Course folder name
    course_folder_name = st.text_input("Enter the course folder name:", "")

    if st.button("Process Document"):
        if not course_folder_name.strip():
            st.error("Please provide a course folder name to proceed.")
            return

        try:
            # Retrieve the top-level folder
            with st.spinner("Looking for the top-level folder..."):
                top_folder_name = "1 WSQ Documents"
                wsq_folder_list = drive_service.files().list(
                    q=f"name='{top_folder_name}' and mimeType='application/vnd.google-apps.folder'",
                    fields="files(id, name)"
                ).execute().get("files", [])

                if not wsq_folder_list:
                    st.error(f"Top-level folder '{top_folder_name}' not found.")
                    return

                wsq_documents_folder_id = wsq_folder_list[0]["id"]

            # Search for the course folder
            with st.spinner("Searching for the course folder..."):
                course_folders = drive_service.files().list(
                    q=f"'{wsq_documents_folder_id}' in parents and mimeType='application/vnd.google-apps.folder'",
                    fields="files(id, name)"
                ).execute().get("files", [])

                matching_course_folder = next(
                    (folder for folder in course_folders if folder["name"].strip().lower() == course_folder_name.strip().lower()),
                    None
                )

                if not matching_course_folder:
                    st.error(f"Course folder '{course_folder_name}' not found.")
                    return

                course_folder_id = matching_course_folder["id"]
                st.success(f"Found course folder: {course_folder_name}")

            # Process the folder
            with st.spinner("Processing the course folder..."):
                result = process_course_folder_direct(course_folder_id, drive_service, abbreviations)
                if not result:
                    st.error("No valid assessment plan or files found in the folder.")
                    return

                assessment_plan = result["assessment_plan"]
                method_data = result["method_data"]

                st.write(f"Assessment Plan: {assessment_plan['name']}")
                st.write(f"Method Data: {json.dumps(method_data, indent=2)}")

                # Download the assessment plan
                with st.spinner("Downloading assessment plan..."):
                    plan_path = download_file(assessment_plan["id"], assessment_plan["name"], drive_service)
                if not plan_path:
                    st.error(f"Failed to download assessment plan: {assessment_plan['name']}")
                    return

                # Download and append Q&A
                with st.spinner("Downloading and appending Q&A documents..."):
                    for abbr, doc_dict in method_data.items():
                        for doc_type in ["question", "answer"]:
                            doc_info = doc_dict.get(doc_type)
                            if doc_info:
                                local_path = download_file(doc_info["id"], doc_info["name"], drive_service)
                                if local_path:
                                    doc_info["local_path"] = local_path

                # Merge Q&A into the annex
                with st.spinner("Merging Q&A into the annex..."):
                    merged_doc_path, changes_made = insert_answers_under_heading(plan_path, heading_map, method_data)
                if changes_made:
                    # with st.spinner("Updating version control..."):
                    #     updated_doc_path = update_cover_page_version(merged_doc_path)
                    #     update_version_control_record(
                    #         doc_path=updated_doc_path,
                    #         changes="Added Assessment Questions & Answers into Annex D",
                    #         developer="Tertiary Infotech"
                    #     )
                    with st.spinner("Renaming and uploading the file..."):
                        final_doc_path = bump_filename_version(merged_doc_path)
                        upload_updated_doc(
                            drive_service=drive_service,
                            file_id=assessment_plan["id"],
                            local_doc_path=final_doc_path,
                            original_filename=assessment_plan["name"]
                        )
                    st.success(f"Updated assessment plan uploaded: {assessment_plan['name']}")
                else:
                    st.info("No changes made to the assessment plan.")

        except Exception as e:
            st.error(f"An error occurred: {e}")
